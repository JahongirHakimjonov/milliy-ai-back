from docx import Document
from docx.shared import Pt

docx_path = "/mnt/data/dasturlash_haqida.docx"
doc = Document()
doc.add_heading("Dasturlash nima?", 0)


# Stil funksiyasi
def add_section(title, text):
    doc.add_heading(title, level=1)
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(12)


# Bo'limlar
add_section(
    "1. Kirish – Dasturlash nima va nima uchun kerak",
    """
Dasturlash – bu kompyuterga aniq vazifalarni bajarishni o‘rgatish san’ati va fanidir. 
U insonning fikrlash jarayonini mantiqiy shaklda ifodalab, kompyuter tushuna oladigan buyruqlarga aylantiradi. 
Har bir dastur – bu kompyuter bajaradigan qadamlar ketma-ketligi, ya’ni algoritmdir. 
Dasturlash bugungi raqamli davrda eng muhim ko‘nikmalardan biri bo‘lib, u orqali inson texnologiyani boshqaradi, yangi mahsulotlar yaratadi va murakkab muammolarga yechim topadi.
""",
)

add_section(
    "2. Dasturlashning qisqacha tarixi",
    """
Dasturlash tarixi XIX asrda Ada Lavleysning hisoblash mashinasi uchun yozgan dasturidan boshlanadi. 
XX asr o‘rtalarida kompyuterlarning rivojlanishi bilan birga dasturlash tillari ham paydo bo‘la boshladi. 
1950-yillarda FORTRAN va COBOL kabi dastlabki tillar ishlab chiqildi. 
1970-1980 yillarda C, Pascal va boshqa yuqori darajali tillar paydo bo‘lib, dasturlashni yanada qulaylashtirdi. 
Bugungi kunda esa Python, JavaScript, Java kabi zamonaviy tillar keng qo‘llaniladi.
""",
)

add_section(
    "3. Dasturlash tillari turlari",
    """
Dasturlash tillari asosan ikki turga bo‘linadi: past darajali va yuqori darajali. 
- Past darajali tillar (masalan, Assembly) kompyuter apparatiga yaqin bo‘lib, tezlik va samaradorlikni ta’minlaydi, ammo yozish qiyinroq. 
- Yuqori darajali tillar esa (masalan, Python, Java) inson tiliga yaqinroq bo‘lib, o‘qish va tushunish osonroq. 
Ular orqali murakkab dasturlarni qisqa vaqt ichida yozish mumkin.
""",
)

add_section(
    "4. Algoritm va mantiqiy fikrlashning ahamiyati",
    """
Har bir dastur algoritmga asoslanadi. Algoritm – bu aniq maqsadga erishish uchun bajariladigan qadamlar ketma-ketligi. 
Dasturchining asosiy vazifasi – muammoni yechish uchun eng samarali algoritmni tuzishdir. 
Shu bois, mantiqiy fikrlash va tahlil qilish ko‘nikmalari dasturlashda katta ahamiyatga ega. 
Har qanday tilni o‘rganishdan oldin algoritmlarni tushunish muhimdir.
""",
)

add_section(
    "5. Dasturlash jarayoni bosqichlari",
    """
Dastur yaratish bir nechta bosqichlardan iborat:
1. Muammoni tahlil qilish va talablarni aniqlash  
2. Algoritm tuzish  
3. Kod yozish  
4. Sinovdan o‘tkazish va xatolarni tuzatish  
5. Dasturga texnik xizmat ko‘rsatish va yangilash  

Bu bosqichlar har bir loyihada takrorlanadi va sifatli dastur yaratishda muhim ahamiyatga ega.
""",
)

add_section(
    "6. Real hayotdagi qo‘llanilishi",
    """
Dasturlash hozirda hayotimizning deyarli har bir jabhasida qo‘llaniladi. 
- Veb dasturlash: veb-saytlar va onlayn xizmatlarni yaratish  
- Mobil ilovalar: smartfonlar uchun foydali ilovalar yaratish  
- Sun’iy intellekt: mashina o‘rganishi va avtomatlashtirish  
- IoT: aqlli qurilmalar va uy tizimlarini boshqarish  

Texnologiya rivojlangani sari dasturchilarga bo‘lgan ehtiyoj yanada ortib bormoqda.
""",
)

add_section(
    "7. Xulosa – Kelajakda dasturlashning roli",
    """
Dasturlash XXI asrda eng talabgir va istiqbolli sohalardan biridir. 
U nafaqat texnologiyani yaratadi, balki inson hayotini yengillashtiradi va yangi imkoniyatlar eshigini ochadi. 
Kelajakda sun’iy intellekt, robototexnika va raqamli iqtisodiyotning markazida aynan dasturchilar turadi. 
Shu sababli dasturlashni o‘rganish – bu nafaqat kasbiy ko‘nikma, balki zamon bilan hamnafas yashash kalitidir.
""",
)

# DOCX faylni saqlash
doc.save(docx_path)

###################################################################################3
###################################################################################3
###################################################################################3
###################################################################################3


from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase import pdfmetrics

# Fayl nomi
output_path = "/mnt/data/dasturlash_haqida.pdf"

# PDF hujjatni yaratish
doc = SimpleDocTemplate(output_path, pagesize=A4, title="Dasturlash haqida maqola")

# Unicode fontni ro'yxatdan o'tkazish (o'zbekcha matn uchun)
pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))

# Stil sozlamalari
styles = getSampleStyleSheet()
styles.add(
    ParagraphStyle(
        name="Heading", fontName="HeiseiMin-W3", fontSize=18, spaceAfter=12, alignment=1
    )
)
styles.add(
    ParagraphStyle(
        name="SubHeading", fontName="HeiseiMin-W3", fontSize=14, spaceAfter=8
    )
)
styles.add(
    ParagraphStyle(name="Body", fontName="HeiseiMin-W3", fontSize=12, leading=18)
)

content = []

# --- Sarlavha ---
content.append(Paragraph("Dasturlash nima?", styles["Heading"]))
content.append(Spacer(1, 20))

# --- 1. Kirish ---
content.append(
    Paragraph("1. Kirish – Dasturlash nima va nima uchun kerak", styles["SubHeading"])
)
content.append(
    Paragraph(
        """
Dasturlash – bu kompyuterga aniq vazifalarni bajarishni o‘rgatish san’ati va fanidir. U insonning fikrlash jarayonini mantiqiy shaklda ifodalab, kompyuter tushuna oladigan buyruqlarga aylantiradi. Har bir dastur – bu kompyuter bajaradigan qadamlar ketma-ketligi, ya’ni algoritmdir. Dasturlash bugungi raqamli davrda eng muhim ko‘nikmalardan biri bo‘lib, u orqali inson texnologiyani boshqaradi, yangi mahsulotlar yaratadi va murakkab muammolarga yechim topadi.
""",
        styles["Body"],
    )
)

# --- 2. Tarix ---
content.append(Spacer(1, 12))
content.append(Paragraph("2. Dasturlashning qisqacha tarixi", styles["SubHeading"]))
content.append(
    Paragraph(
        """
Dasturlash tarixi XIX asrda Ada Lavleysning hisoblash mashinasi uchun yozgan dasturidan boshlanadi. XX asr o‘rtalarida kompyuterlarning rivojlanishi bilan birga dasturlash tillari ham paydo bo‘la boshladi. 1950-yillarda FORTRAN va COBOL kabi dastlabki tillar ishlab chiqildi. 1970-1980 yillarda C, Pascal va boshqa yuqori darajali tillar paydo bo‘lib, dasturlashni yanada qulaylashtirdi. Bugungi kunda esa Python, JavaScript, Java kabi zamonaviy tillar keng qo‘llaniladi.
""",
        styles["Body"],
    )
)

# --- 3. Tillarning turlari ---
content.append(Spacer(1, 12))
content.append(Paragraph("3. Dasturlash tillari turlari", styles["SubHeading"]))
content.append(
    Paragraph(
        """
Dasturlash tillari asosan ikki turga bo‘linadi: past darajali va yuqori darajali. 
- **Past darajali tillar** (masalan, Assembly) kompyuter apparatiga yaqin bo‘lib, tezlik va samaradorlikni ta’minlaydi, ammo yozish qiyinroq. 
- **Yuqori darajali tillar** esa (masalan, Python, Java) inson tiliga yaqinroq bo‘lib, o‘qish va tushunish osonroq. Ular orqali murakkab dasturlarni qisqa vaqt ichida yozish mumkin.
""",
        styles["Body"],
    )
)

# --- 4. Algoritm va mantiq ---
content.append(Spacer(1, 12))
content.append(
    Paragraph("4. Algoritm va mantiqiy fikrlashning ahamiyati", styles["SubHeading"])
)
content.append(
    Paragraph(
        """
Har bir dastur algoritmga asoslanadi. Algoritm – bu aniq maqsadga erishish uchun bajariladigan qadamlar ketma-ketligi. Dasturchining asosiy vazifasi – muammoni yechish uchun eng samarali algoritmni tuzishdir. Shu bois, mantiqiy fikrlash va tahlil qilish ko‘nikmalari dasturlashda katta ahamiyatga ega. Har qanday tilni o‘rganishdan oldin algoritmlarni tushunish muhimdir.
""",
        styles["Body"],
    )
)

# --- 5. Jarayon bosqichlari ---
content.append(Spacer(1, 12))
content.append(Paragraph("5. Dasturlash jarayoni bosqichlari", styles["SubHeading"]))
content.append(
    Paragraph(
        """
Dastur yaratish bir nechta bosqichlardan iborat:
1. Muammoni tahlil qilish va talablarni aniqlash  
2. Algoritm tuzish  
3. Kod yozish  
4. Sinovdan o‘tkazish va xatolarni tuzatish  
5. Dasturga texnik xizmat ko‘rsatish va yangilash  

Bu bosqichlar har bir loyihada takrorlanadi va sifatli dastur yaratishda muhim ahamiyatga ega.
""",
        styles["Body"],
    )
)

# --- 6. Qo'llanilishi ---
content.append(Spacer(1, 12))
content.append(Paragraph("6. Real hayotdagi qo‘llanilishi", styles["SubHeading"]))
content.append(
    Paragraph(
        """
Dasturlash hozirda hayotimizning deyarli har bir jabhasida qo‘llaniladi. 
- **Veb dasturlash:** veb-saytlar va onlayn xizmatlarni yaratish  
- **Mobil ilovalar:** smartfonlar uchun foydali ilovalar yaratish  
- **Sun’iy intellekt:** mashina o‘rganishi va avtomatlashtirish  
- **IoT:** aqlli qurilmalar va uy tizimlarini boshqarish  

Texnologiya rivojlangani sari dasturchilarga bo‘lgan ehtiyoj yanada ortib bormoqda.
""",
        styles["Body"],
    )
)

content.append(Spacer(1, 12))
content.append(
    Paragraph("7. Xulosa – Kelajakda dasturlashning roli", styles["SubHeading"])
)
content.append(
    Paragraph(
        """
Dasturlash XXI asrda eng talabgir va istiqbolli sohalardan biridir. U nafaqat texnologiyani yaratadi, balki inson hayotini yengillashtiradi va yangi imkoniyatlar eshigini ochadi. Kelajakda sun’iy intellekt, robototexnika va raqamli iqtisodiyotning markazida aynan dasturchilar turadi. Shu sababli dasturlashni o‘rganish – bu nafaqat kasbiy ko‘nikma, balki zamon bilan hamnafas yashash kalitidir.
""",
        styles["Body"],
    )
)

# PDFni saqlash
doc.build(content)
