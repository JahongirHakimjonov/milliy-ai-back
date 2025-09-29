import os
import re
import uuid
from textwrap import wrap
from typing import List

from asgiref.sync import async_to_sync
from django.core.files import File
from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from apps.chat.enums.action import FileFormat
from apps.chat.models.chat import ChatResource
from apps.chat.services.ai import AIService
from apps.users.models.users import User

INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)", flags=re.DOTALL)


def add_runs_with_inline_format(paragraph, text: str):
    parts = INLINE_PATTERN.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            try:
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            except Exception:
                pass
        else:
            paragraph.add_run(part)


def save_markdown_to_docx(md_text: str, filename: str = "output.docx"):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.splitlines()
    in_code_block = False
    code_block_lines: List[str] = []

    for raw in lines:
        line = raw.rstrip()

        # Code block handling
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_block_lines = []
                continue
            else:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_block_lines))
                r.font.name = "Courier New"
                r.font.size = Pt(9)
                in_code_block = False
                code_block_lines = []
                continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        # Headings
        if line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 2")
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 3")
            continue
        if line.startswith("##### "):
            doc.add_paragraph(line[5:].strip(), style="Heading 4")
            continue

        # Unordered lists
        m_un = re.match(r"^(\s*)([-\*])\s+(.*)$", line)
        if m_un:
            indent = len(m_un.group(1))
            text = m_un.group(3)
            p = doc.add_paragraph(style="List Bullet")
            if indent >= 4:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 4))
            add_runs_with_inline_format(p, text)
            continue

        # Ordered lists
        m_ord = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_ord:
            indent = len(m_ord.group(1))
            text = m_ord.group(3)
            p = doc.add_paragraph(style="List Number")
            if indent >= 4:
                p.paragraph_format.left_indent = Inches(0.25 * (indent // 4))
            add_runs_with_inline_format(p, text)
            continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_runs_with_inline_format(p, line[2:].strip())
            continue

        # Paragraph
        p = doc.add_paragraph()
        add_runs_with_inline_format(p, line)

    for p in doc.paragraphs:
        p.paragraph_format.space_after = Pt(6)

    doc.save(filename)


def save_markdown_to_pdf_reportlab(md_text: str, filename: str = "output.pdf"):
    """Convert markdown-like text into a simple but structured PDF."""
    try:
        pdfmetrics.registerFont(TTFont("CourierNew", "Courier_New.ttf"))
        monospace_name = "CourierNew"
    except Exception:
        monospace_name = "Courier"

    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4
    left_margin = 50
    y = height - 50

    lines = md_text.splitlines()
    in_code_block = False
    code_block_collect = []

    for raw in lines:
        line = raw.rstrip()

        # Code block handling
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_block_collect = []
                continue
            else:
                for code_line in code_block_collect:
                    for wrapped in wrap(code_line, 95):
                        c.setFont(monospace_name, 9)
                        c.drawString(left_margin + 10, y, wrapped)
                        y -= 14
                        if y < 50:
                            c.showPage()
                            y = height - 50
                y -= 8
                in_code_block = False
                continue

        if in_code_block:
            code_block_collect.append(line)
            continue

        if not line.strip():
            y -= 8
            if y < 50:
                c.showPage()
                y = height - 50
            continue

        # Headings
        if line.startswith("# "):
            c.setFont("Helvetica-Bold", 16)
            c.drawString(left_margin, y, line[2:].strip())
            y -= 20
            continue
        if line.startswith("## "):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(left_margin, y, line[3:].strip())
            y -= 18
            continue
        if line.startswith("### "):
            c.setFont("Helvetica-Bold", 12)
            c.drawString(left_margin, y, line[4:].strip())
            y -= 16
            continue
        if line.startswith("##### "):
            c.setFont("Helvetica-Bold", 11)
            c.drawString(left_margin, y, line[5:].strip())
            y -= 14
            continue

        # Unordered list
        m_un = re.match(r"^(\s*)([-\*])\s+(.*)$", line)
        if m_un:
            indent = len(m_un.group(1))
            text = m_un.group(3)
            bullet_x = left_margin + (indent // 4) * 14
            c.setFont("Helvetica", 11)
            c.drawString(bullet_x, y, "•")
            for wrapped in wrap(text, 80 - (indent // 4) * 4):
                c.drawString(bullet_x + 14, y, wrapped)
                y -= 14
                if y < 50:
                    c.showPage()
                    y = height - 50
            y -= 4
            continue

        # Ordered list
        m_ord = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_ord:
            indent = len(m_ord.group(1))
            num = m_ord.group(2)
            text = m_ord.group(3)
            bullet_x = left_margin + (indent // 4) * 14
            c.setFont("Helvetica", 11)
            c.drawString(bullet_x, y, f"{num}.")
            for wrapped in wrap(text, 80 - (indent // 4) * 4):
                c.drawString(bullet_x + 20, y, wrapped)
                y -= 14
                if y < 50:
                    c.showPage()
                    y = height - 50
            y -= 4
            continue

        # Plain text
        text = re.sub(r"\*\*(.+?)\*\*", lambda m: m.group(1), line)
        text = re.sub(r"\*(.+?)\*", lambda m: m.group(1), text)
        text = re.sub(r"`(.+?)`", lambda m: m.group(1), text)

        for wrapped in wrap(text, 95):
            c.setFont("Helvetica", 11)
            c.drawString(left_margin, y, wrapped)
            y -= 14
            if y < 50:
                c.showPage()
                y = height - 50

    c.save()


# ---------------------- DB SAVE ----------------------
def save_chat_resource(user: User, filepath: str) -> ChatResource:
    """Save a generated file to ChatResource for a user."""
    if not user or not isinstance(user, User):
        raise ValueError("❌ Invalid user provided.")
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"❌ File not found: {filepath}")

    filename = os.path.basename(filepath)
    with open(filepath, "rb") as f:
        resource = ChatResource.objects.create(user=user, file=File(f, name=filename))
        file_id = async_to_sync(AIService().create_file)(file=f)
        resource.file_id = file_id
        resource.save(update_fields=["file_id"])
    return resource


def file_service(text: str, file_format: str, user: User) -> List[int]:
    """
    Generate a DOCX or PDF from markdown text and save it as a ChatResource.
    Returns: ChatResource instance.
    """
    if file_format.lower() not in {FileFormat.PDF, FileFormat.DOCX}:
        raise ValueError("❌ Unsupported format. Use 'pdf' or 'docx'.")

    filename = f"generated_{uuid.uuid4().hex[:8]}.{file_format.lower()}"

    if file_format.lower() == FileFormat.PDF:
        save_markdown_to_pdf_reportlab(text, filename)
    else:
        save_markdown_to_docx(text, filename)

    # Save the file in the database
    resource = save_chat_resource(user, filename)
    file_ids = [resource.id]

    try:
        os.remove(filename)
    except OSError:
        pass

    return file_ids
